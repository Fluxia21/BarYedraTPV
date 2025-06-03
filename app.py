import os
import logging
import json
from datetime import datetime, date, timedelta
from decimal import Decimal
from io import BytesIO
import openpyxl
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, session
from functools import wraps
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import DeclarativeBase
from werkzeug.middleware.proxy_fix import ProxyFix

# Configure logging
logging.basicConfig(level=logging.DEBUG)

class Base(DeclarativeBase):
    pass

db = SQLAlchemy(model_class=Base)

# Create the app
app = Flask(__name__)
app.secret_key = "yedra-bar-secret-key-2024"
# ProxyFix no necesario para uso local
# app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

# Configure PostgreSQL database
app.config["SQLALCHEMY_DATABASE_URI"] = os.environ.get("DATABASE_URL")
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    "pool_recycle": 300,
    "pool_pre_ping": True,
}

# Initialize the app with the extension
db.init_app(app)

# Import models after db initialization
from models import Mesa, Producto, Pedido, Ticket, MovimientoStock, Empleado, Fichaje

# Add custom Jinja2 filter for JSON parsing
# Login configuration
LOGIN_USERNAME = "JAVI"
LOGIN_PIN = "5555"

# Login required decorator
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('logged_in'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

@app.template_filter('from_json')
def from_json_filter(value):
    """Convert JSON string to Python object"""
    if value:
        try:
            return json.loads(value)
        except (json.JSONDecodeError, TypeError):
            return []
    return []

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page with username and PIN authentication"""
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        pin = request.form.get('pin', '').strip()
        
        if username == LOGIN_USERNAME and pin == LOGIN_PIN:
            session['logged_in'] = True
            session['username'] = username
            flash('Acceso autorizado', 'success')
            return redirect(url_for('index'))
        else:
            flash('Usuario o PIN incorrecto', 'error')
            return render_template('login.html', error=True)
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    """Logout and clear session"""
    session.clear()
    flash('Sesión cerrada', 'info')
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    """Página principal - redirigir a terraza"""
    return redirect(url_for('terraza'))

@app.route('/terraza')
@login_required
def terraza():
    """Página de mesas de terraza"""
    mesas = Mesa.query.filter_by(zona='Terraza').order_by(Mesa.numero).all()
    return render_template('terraza.html', mesas=mesas)

@app.route('/sala')
@login_required
def sala():
    """Página de mesas de sala"""
    mesas = Mesa.query.filter_by(zona='Sala').order_by(Mesa.numero).all()
    return render_template('sala.html', mesas=mesas)

@app.route('/barra')
@login_required
def barra():
    """Página de mesas de barra"""
    mesas = Mesa.query.filter_by(zona='Barra').order_by(Mesa.numero).all()
    return render_template('barra.html', mesas=mesas)

def get_menu_recommendations(mesa_id, limit=5):
    """Generate personalized menu recommendations based on previous orders"""
    # Obtener historial de pedidos de esta mesa (últimos 30 días)
    from datetime import datetime, timedelta
    fecha_limite = datetime.utcnow() - timedelta(days=30)
    
    # Contar productos más pedidos en esta mesa
    pedidos_mesa = db.session.query(Pedido).filter(
        Pedido.mesa_id == mesa_id,
        Pedido.fecha_creacion >= fecha_limite,
        Pedido.estado == 'pagado'
    ).all()
    
    producto_frecuencia = {}
    for pedido in pedidos_mesa:
        if pedido.productos:
            try:
                productos_pedido = json.loads(pedido.productos)
                for producto_id, cantidad in productos_pedido.items():
                    producto_id = int(producto_id)
                    if producto_id not in producto_frecuencia:
                        producto_frecuencia[producto_id] = 0
                    producto_frecuencia[producto_id] += cantidad
            except (json.JSONDecodeError, ValueError):
                continue
    
    # Si no hay historial en esta mesa, obtener productos populares generales
    if not producto_frecuencia:
        pedidos_generales = db.session.query(Pedido).filter(
            Pedido.fecha_creacion >= fecha_limite,
            Pedido.estado == 'pagado'
        ).limit(50).all()
        
        for pedido in pedidos_generales:
            if pedido.productos:
                try:
                    productos_pedido = json.loads(pedido.productos)
                    for producto_id, cantidad in productos_pedido.items():
                        producto_id = int(producto_id)
                        if producto_id not in producto_frecuencia:
                            producto_frecuencia[producto_id] = 0
                        producto_frecuencia[producto_id] += cantidad
                except (json.JSONDecodeError, ValueError):
                    continue
    
    # Ordenar por frecuencia y obtener productos
    productos_recomendados = []
    if producto_frecuencia:
        productos_ordenados = sorted(producto_frecuencia.items(), key=lambda x: x[1], reverse=True)[:limit]
        for producto_id, frecuencia in productos_ordenados:
            producto = Producto.query.get(producto_id)
            if producto and producto.activo:
                productos_recomendados.append({
                    'producto': producto,
                    'frecuencia': frecuencia
                })
    
    return productos_recomendados

@app.route('/mesa/<int:mesa_id>')
@login_required
def table_detail(mesa_id):
    """Table detail view for managing orders"""
    mesa = Mesa.query.get_or_404(mesa_id)
    productos = Producto.query.filter_by(activo=True).order_by(Producto.categoria, Producto.orden, Producto.nombre).all()
    
    # Verificar si es mesa de terraza para aplicar suplemento
    es_terraza = mesa.zona.lower() == 'terraza'
    suplemento_terraza = Decimal('0.20') if es_terraza else Decimal('0.00')
    
    # Agrupar productos por categoría y preparar precios con suplemento
    productos_por_categoria = {}
    productos_dict = {}
    productos_dict_con_precios = {}
    
    for producto in productos:
        if producto.categoria not in productos_por_categoria:
            productos_por_categoria[producto.categoria] = []
        productos_por_categoria[producto.categoria].append(producto)
        productos_dict[producto.id] = producto
        
        # Crear una versión del producto con precios calculados para terraza
        precio_base = float(producto.precio)
        # No aplicar suplemento si el producto tiene sin_suplemento_terraza = True
        aplicar_suplemento = es_terraza and not getattr(producto, 'sin_suplemento_terraza', False)
        precio_final = precio_base + 0.20 if aplicar_suplemento else precio_base
        

        
        productos_dict_con_precios[producto.id] = {
            'id': producto.id,
            'nombre': producto.nombre,
            'precio_base': precio_base,
            'precio_final': precio_final,
            'categoria': producto.categoria,
            'descripcion': producto.descripcion,
            'activo': producto.activo
        }
    
    # Get current order for this table if exists
    pedido_actual = Pedido.query.filter_by(mesa_id=mesa_id, estado='abierto').first()
    
    # Obtener recomendaciones personalizadas
    recomendaciones = get_menu_recommendations(mesa_id)
    
    return render_template('table_detail.html', 
                         mesa=mesa, 
                         productos_por_categoria=productos_por_categoria,
                         productos_dict=productos_dict,
                         productos_dict_con_precios=productos_dict_con_precios,
                         pedido_actual=pedido_actual,
                         es_terraza=es_terraza,
                         suplemento_terraza=suplemento_terraza,
                         recomendaciones=recomendaciones)

@app.route('/add_to_order', methods=['POST'])
@login_required
def add_to_order():
    """Add product to table order - accumulates by rounds"""
    mesa_id = int(request.form.get('mesa_id'))
    producto_id = int(request.form.get('producto_id'))
    cantidad = int(request.form.get('cantidad', 1))
    
    mesa = Mesa.query.get_or_404(mesa_id)
    producto = Producto.query.get_or_404(producto_id)
    
    # Aplicar suplemento de terraza si corresponde
    precio_producto = float(producto.precio)
    if mesa.zona.lower() == 'terraza':
        precio_producto += 0.20
    
    # Get or create current order
    pedido = Pedido.query.filter_by(mesa_id=mesa_id, estado='abierto').first()
    if not pedido:
        pedido = Pedido(
            mesa_id=mesa_id, 
            productos='{}',
            total=0.0, 
            estado='abierto'
        )
        db.session.add(pedido)
        mesa.estado = 'ocupada'
        db.session.flush()  # Get the ID
    
    # Update products (storing as JSON with product_id as key and quantity as value)
    import json
    if pedido.productos:
        productos_dict = json.loads(pedido.productos)
    else:
        productos_dict = {}
    
    # Add or update product quantity (accumulate by rounds)
    producto_id_str = str(producto_id)
    if producto_id_str in productos_dict:
        productos_dict[producto_id_str] += cantidad
    else:
        productos_dict[producto_id_str] = cantidad
    
    # Save updated products and calculate total
    pedido.productos = json.dumps(productos_dict)
    
    # Calculate total with terraza supplement if applicable
    total = 0
    for pid, qty in productos_dict.items():
        prod = Producto.query.get(int(pid))
        if prod:
            precio_unitario = float(prod.precio)
            if mesa.zona.lower() == 'terraza' and not getattr(prod, 'sin_suplemento_terraza', False):
                precio_unitario += 0.20
            total += precio_unitario * qty
    
    pedido.total = total
    
    db.session.commit()
    
    # Redirect back to table detail page
    return redirect(url_for('table_detail', mesa_id=mesa_id))

@app.route('/update_quantity', methods=['POST'])
def update_quantity():
    """Update product quantity in current order"""
    mesa_id = int(request.form.get('mesa_id'))
    producto_id = int(request.form.get('producto_id'))
    nueva_cantidad = int(request.form.get('cantidad'))
    
    pedido = Pedido.query.filter_by(mesa_id=mesa_id, estado='abierto').first()
    if not pedido:
        return '', 404
    
    import json
    productos_dict = json.loads(pedido.productos) if pedido.productos else {}
    producto_id_str = str(producto_id)
    
    if producto_id_str in productos_dict:
        if nueva_cantidad <= 0:
            del productos_dict[producto_id_str]
        else:
            productos_dict[producto_id_str] = nueva_cantidad
    
    # Recalculate total with terraza supplement if applicable
    mesa = Mesa.query.get(mesa_id)
    total = 0
    for pid, qty in productos_dict.items():
        prod = Producto.query.get(int(pid))
        if prod:
            precio_unitario = float(prod.precio)
            if mesa.zona.lower() == 'terraza' and not getattr(prod, 'sin_suplemento_terraza', False):
                precio_unitario += 0.20
            total += precio_unitario * qty
    
    pedido.productos = json.dumps(productos_dict)
    pedido.total = total
    
    # If no products left, delete the order
    if not productos_dict:
        mesa = Mesa.query.get(mesa_id)
        mesa.estado = 'libre'
        db.session.delete(pedido)
    
    db.session.commit()
    return redirect(url_for('table_detail', mesa_id=mesa_id))

@app.route('/close_order/<int:mesa_id>', methods=['GET', 'POST'])
def close_order(mesa_id):
    """Close order and set table as ready to pay"""
    mesa = Mesa.query.get_or_404(mesa_id)
    pedido = Pedido.query.filter_by(mesa_id=mesa_id, estado='abierto').first()
    
    if pedido:
        pedido.estado = 'cerrado'
        mesa.estado = 'pagada'
        db.session.commit()
        flash('Pedido cerrado. Selecciona método de pago.', 'success')
        # Redirect to table detail to show payment options
        return redirect(url_for('table_detail', mesa_id=mesa_id))
    
    return redirect(url_for('index'))

@app.route('/pay_order/<int:mesa_id>', methods=['GET', 'POST'])
def pay_order(mesa_id):
    """Process payment for table with method selection"""
    if request.method == 'POST':
        payment_method = request.form.get('forma_pago', 'efectivo')
    else:
        payment_method = request.args.get('payment_method', 'efectivo')
    mesa = Mesa.query.get_or_404(mesa_id)
    
    # Look for closed or open order to pay
    pedido = Pedido.query.filter_by(mesa_id=mesa_id).filter(
        Pedido.estado.in_(['cerrado', 'abierto'])
    ).first()
    
    if not pedido:
        flash('No hay pedido para pagar en esta mesa', 'error')
        return redirect(url_for('index'))
    
    # Process payment
    pedido.forma_pago = payment_method
    pedido.fecha_pago = datetime.utcnow()
    pedido.estado = 'pagado'
    
    # Create ticket for automatic printing
    ticket = Ticket(
        pedido_id=pedido.id
    )
    db.session.add(ticket)
    
    # Free the table
    mesa.estado = 'libre'
    
    db.session.commit()
    flash(f'Pago procesado ({payment_method}). Ticket generado automáticamente.', 'success')
    
    # Automatically redirect to print receipt
    return redirect(url_for('print_receipt', pedido_id=pedido.id))

@app.route('/print_receipt/<int:pedido_id>')
def print_receipt(pedido_id):
    """Generate printable receipt with detailed product information"""
    pedido = Pedido.query.get_or_404(pedido_id)
    mesa = Mesa.query.get(pedido.mesa_id)
    
    # Parse products and get detailed information
    import json
    productos_dict = json.loads(pedido.productos) if pedido.productos else {}
    
    # Build detailed products list with quantities and prices
    productos_detalle = []
    for producto_id_str, cantidad in productos_dict.items():
        producto = Producto.query.get(int(producto_id_str))
        if producto:
            productos_detalle.append({
                'nombre': producto.nombre,
                'precio': float(producto.precio),
                'cantidad': cantidad,
                'subtotal': float(producto.precio) * cantidad
            })
    
    return render_template('receipt_print.html', 
                         pedido=pedido, 
                         mesa=mesa,
                         productos_detalle=productos_detalle)

@app.route('/cash_register')
def cash_register():
    """Daily cash register view"""
    today = date.today()
    
    # Get today's paid orders
    pedidos_hoy = Pedido.query.filter(
        db.func.date(Pedido.fecha_pago) == today,
        Pedido.estado == 'pagado'
    ).all()
    
    total_efectivo = sum(p.total for p in pedidos_hoy if p.forma_pago == 'efectivo')
    total_tarjeta = sum(p.total for p in pedidos_hoy if p.forma_pago == 'tarjeta')
    total_dia = total_efectivo + total_tarjeta
    
    return render_template('cash_register.html',
                         pedidos_hoy=pedidos_hoy,
                         total_efectivo=total_efectivo,
                         total_tarjeta=total_tarjeta,
                         total_dia=total_dia,
                         fecha=today)

@app.route('/close_cash_register', methods=['POST'])
def close_cash_register():
    """Close daily cash register"""
    today = date.today()
    
    # Archive today's data (in a real system, you might move to archive table)
    pedidos_hoy = Pedido.query.filter(
        db.func.date(Pedido.fecha_pago) == today,
        Pedido.estado == 'pagado'
    ).all()
    
    if pedidos_hoy:
        flash(f'Caja cerrada. {len(pedidos_hoy)} pedidos archivados.', 'success')
    else:
        flash('No hay pedidos para cerrar hoy.', 'info')
    
    return redirect(url_for('cash_register'))

@app.route('/products')
@login_required
def products():
    """Product management view"""
    productos = Producto.query.filter_by(activo=True).order_by(Producto.categoria, Producto.nombre).all()
    
    # Agrupar productos por categoría (ya ordenados alfabéticamente)
    productos_por_categoria = {}
    for producto in productos:
        if producto.categoria not in productos_por_categoria:
            productos_por_categoria[producto.categoria] = []
        productos_por_categoria[producto.categoria].append(producto)
    
    return render_template('products.html', productos_por_categoria=productos_por_categoria)

@app.route('/add_product', methods=['POST'])
def add_product():
    """Add new product with photo, category and description"""
    nombre = request.form.get('nombre')
    precio_str = request.form.get('precio', '').strip().lower()
    if precio_str in ['nan', 'infinity', '-infinity', 'inf', '-inf']:
        flash('Precio no válido', 'error')
        return redirect(url_for('products'))
    try:
        precio = float(request.form.get('precio'))
        if precio < 0:
            flash('El precio debe ser positivo', 'error')
            return redirect(url_for('products'))
    except (ValueError, TypeError):
        flash('Precio no válido', 'error')
        return redirect(url_for('products'))
    categoria = request.form.get('categoria', 'General')
    descripcion = request.form.get('descripcion', '')
    
    # Manejar carga de imagen si se proporcionó
    foto_url = ''
    if 'imagen_producto' in request.files:
        file = request.files['imagen_producto']
        if file and file.filename != '':
            # Validar tipo de archivo
            allowed_extensions = {'png', 'jpg', 'jpeg', 'gif'}
            if '.' in file.filename and file.filename.rsplit('.', 1)[1].lower() in allowed_extensions:
                # Generar nombre único para el archivo
                import uuid
                import os
                filename = str(uuid.uuid4()) + '.' + file.filename.rsplit('.', 1)[1].lower()
                filepath = os.path.join('static', 'uploads', 'products', filename)
                
                # Crear directorio si no existe
                os.makedirs(os.path.dirname(filepath), exist_ok=True)
                
                # Guardar archivo
                file.save(filepath)
                foto_url = f'/static/uploads/products/{filename}'
            else:
                flash('Tipo de archivo no válido. Solo se permiten PNG, JPG, JPEG, GIF', 'error')
                return redirect(url_for('products'))
    
    producto = Producto(
        nombre=nombre, 
        precio=precio,
        categoria=categoria,
        foto_url=foto_url,
        descripcion=descripcion
    )
    db.session.add(producto)
    db.session.commit()
    
    flash(f'Producto {nombre} añadido correctamente en {categoria}', 'success')
    return redirect(url_for('products'))

@app.route('/edit_product/<int:producto_id>')
def edit_product(producto_id):
    """Show edit product form"""
    producto = Producto.query.get_or_404(producto_id)
    
    # Obtener todas las categorías únicas para el select
    categorias = db.session.query(Producto.categoria).distinct().all()
    categorias = [cat[0] for cat in categorias if cat[0]]
    
    return render_template('edit_product.html', producto=producto, categorias=categorias)

@app.route('/update_product/<int:producto_id>', methods=['POST'])
def update_product(producto_id):
    """Update product information"""
    producto = Producto.query.get_or_404(producto_id)
    
    # Validar datos del formulario
    nombre = request.form.get('nombre', '').strip()
    if not nombre:
        flash('El nombre del producto es obligatorio', 'error')
        return redirect(url_for('edit_product', producto_id=producto_id))
    
    precio_str = request.form.get('precio', '').strip().lower()
    if precio_str in ['nan', 'infinity', '-infinity', 'inf', '-inf']:
        flash('Precio no válido', 'error')
        return redirect(url_for('edit_product', producto_id=producto_id))
    try:
        precio = float(request.form.get('precio'))
        if precio < 0:
            flash('El precio debe ser positivo', 'error')
            return redirect(url_for('edit_product', producto_id=producto_id))
    except (ValueError, TypeError):
        flash('Precio no válido', 'error')
        return redirect(url_for('edit_product', producto_id=producto_id))
    
    # Manejar carga de imagen si se proporcionó
    foto_url = producto.foto_url  # Mantener la imagen existente por defecto
    if 'imagen_producto' in request.files:
        file = request.files['imagen_producto']
        if file and file.filename != '':
            # Validar tipo de archivo
            allowed_extensions = {'png', 'jpg', 'jpeg', 'gif'}
            if '.' in file.filename and file.filename.rsplit('.', 1)[1].lower() in allowed_extensions:
                # Generar nombre único para el archivo
                import uuid
                import os
                filename = str(uuid.uuid4()) + '.' + file.filename.rsplit('.', 1)[1].lower()
                filepath = os.path.join('static', 'uploads', 'products', filename)
                
                # Crear directorio si no existe
                os.makedirs(os.path.dirname(filepath), exist_ok=True)
                
                # Guardar archivo
                file.save(filepath)
                foto_url = f'/static/uploads/products/{filename}'
            else:
                flash('Tipo de archivo no válido. Solo se permiten PNG, JPG, JPEG, GIF', 'error')
                return redirect(url_for('edit_product', producto_id=producto_id))
    
    # Actualizar los campos del producto
    producto.nombre = nombre
    producto.precio = precio
    producto.categoria = request.form.get('categoria', 'General')
    producto.foto_url = foto_url
    producto.descripcion = request.form.get('descripcion', '')
    
    db.session.commit()
    
    flash(f'Producto {nombre} actualizado correctamente', 'success')
    return redirect(url_for('products'))

@app.route('/delete_product/<int:producto_id>')
def delete_product(producto_id):
    """Delete product"""
    producto = Producto.query.get_or_404(producto_id)
    db.session.delete(producto)
    db.session.commit()
    
    flash(f'Producto {producto.nombre} eliminado', 'success')
    return redirect(url_for('products'))

@app.route('/inventory')
def inventory():
    """Inventory management view"""
    # Orden inteligente de categorías para el flujo de trabajo del bar
    orden_categorias = {
        'Desayunos': 1,     # Primeros del día
        'Bocadillos': 2,    # Comida principal
        'Raciones': 3,      # Platos compartidos
        'Cervezas': 4,      # Bebidas principales
        'Vinos': 5,         # Bebidas con comida
        'Refrescos': 6,     # Bebidas sin alcohol
        'Combinados': 7,    # Cocktails y mezclas
        'Licores': 8,       # Bebidas destiladas
        'Helados': 9        # Postres
    }
    
    # Obtener todos los productos activos
    productos_db = Producto.query.filter_by(activo=True).all()
    
    # Ordenar por categoría inteligente y luego por nombre
    productos = sorted(productos_db, key=lambda p: (
        orden_categorias.get(p.categoria, 999),  # 999 para categorías no definidas
        p.nombre
    ))
    
    # Estadísticas de stock
    productos_agotados = [p for p in productos if p.stock_agotado]
    productos_bajo_stock = [p for p in productos if p.stock_bajo and not p.stock_agotado]
    
    # Agrupar productos por categoría para mejor visualización
    productos_por_categoria = {}
    for producto in productos:
        categoria = producto.categoria
        if categoria not in productos_por_categoria:
            productos_por_categoria[categoria] = []
        productos_por_categoria[categoria].append(producto)
    
    return render_template('inventory.html', 
                         productos=productos,
                         productos_por_categoria=productos_por_categoria,
                         productos_agotados=productos_agotados,
                         productos_bajo_stock=productos_bajo_stock)

@app.route('/update_stock', methods=['POST'])
def update_stock():
    """Update product stock"""
    producto_id = int(request.form.get('producto_id'))
    nuevo_stock = int(request.form.get('nuevo_stock'))
    motivo = request.form.get('motivo', 'ajuste_inventario')
    
    producto = Producto.query.get_or_404(producto_id)
    stock_anterior = producto.stock_actual
    
    # Crear movimiento de stock
    movimiento = MovimientoStock(
        producto_id=producto_id,
        tipo_movimiento='ajuste',
        cantidad=nuevo_stock - stock_anterior,
        motivo=motivo,
        stock_anterior=stock_anterior,
        stock_nuevo=nuevo_stock
    )
    
    # Actualizar stock
    producto.stock_actual = nuevo_stock
    producto.fecha_actualizacion = datetime.utcnow()
    
    db.session.add(movimiento)
    db.session.commit()
    
    flash(f'Stock de {producto.nombre} actualizado a {nuevo_stock}', 'success')
    return redirect(url_for('inventory'))

@app.route('/add_stock', methods=['POST'])
def add_stock():
    """Add stock to product"""
    producto_id = int(request.form.get('producto_id'))
    cantidad = int(request.form.get('cantidad'))
    motivo = request.form.get('motivo', 'compra')
    
    producto = Producto.query.get_or_404(producto_id)
    stock_anterior = producto.stock_actual
    nuevo_stock = stock_anterior + cantidad
    
    # Crear movimiento de entrada
    movimiento = MovimientoStock(
        producto_id=producto_id,
        tipo_movimiento='entrada',
        cantidad=cantidad,
        motivo=motivo,
        stock_anterior=stock_anterior,
        stock_nuevo=nuevo_stock
    )
    
    # Actualizar stock
    producto.stock_actual = nuevo_stock
    producto.fecha_actualizacion = datetime.utcnow()
    
    db.session.add(movimiento)
    db.session.commit()
    
    flash(f'Añadidas {cantidad} unidades a {producto.nombre}', 'success')
    return redirect(url_for('inventory'))

@app.route('/reports')
def reports():
    """Sales reports dashboard"""
    today = date.today()
    inicio_semana = today - timedelta(days=today.weekday())
    inicio_mes = today.replace(day=1)
    
    # Calcular estadísticas rápidas
    pedidos_hoy = Pedido.query.filter(
        db.func.date(Pedido.fecha_pago) == today,
        Pedido.estado == 'pagado'
    ).count()
    
    ventas_hoy = db.session.query(db.func.sum(Pedido.total)).filter(
        db.func.date(Pedido.fecha_pago) == today,
        Pedido.estado == 'pagado'
    ).scalar() or 0
    
    ventas_semana = db.session.query(db.func.sum(Pedido.total)).filter(
        Pedido.fecha_pago >= inicio_semana,
        Pedido.estado == 'pagado'
    ).scalar() or 0
    
    ventas_mes = db.session.query(db.func.sum(Pedido.total)).filter(
        Pedido.fecha_pago >= inicio_mes,
        Pedido.estado == 'pagado'
    ).scalar() or 0
    
    return render_template('reports.html',
                         pedidos_hoy=pedidos_hoy,
                         ventas_hoy=ventas_hoy,
                         ventas_semana=ventas_semana,
                         ventas_mes=ventas_mes,
                         fecha=today)

@app.route('/export_excel/<periodo>')
def export_excel(periodo):
    """Export sales report to Excel"""
    today = date.today()
    
    if periodo == 'dia':
        fecha_inicio = today
        fecha_fin = today
        titulo = f"Ventas del Día - {today.strftime('%d/%m/%Y')}"
    elif periodo == 'semana':
        fecha_inicio = today - timedelta(days=today.weekday())
        fecha_fin = today
        titulo = f"Ventas de la Semana - {fecha_inicio.strftime('%d/%m')} al {fecha_fin.strftime('%d/%m/%Y')}"
    else:  # mes
        fecha_inicio = today.replace(day=1)
        fecha_fin = today
        titulo = f"Ventas del Mes - {fecha_inicio.strftime('%B %Y')}"
    
    # Obtener datos
    pedidos = Pedido.query.filter(
        Pedido.fecha_pago >= fecha_inicio,
        Pedido.fecha_pago <= fecha_fin + timedelta(days=1),
        Pedido.estado == 'pagado'
    ).all()
    
    # Crear workbook
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Reporte de Ventas"
    
    # Encabezados con estilo
    encabezados = ['Fecha', 'Hora', 'Mesa', 'Total', 'Forma de Pago', 'Ticket #']
    for col, encabezado in enumerate(encabezados, 1):
        celda = ws.cell(row=1, column=col, value=encabezado)
        celda.font = Font(bold=True, color='FFFFFF')
        celda.fill = PatternFill(start_color='2E7D32', end_color='2E7D32', fill_type='solid')
        celda.alignment = Alignment(horizontal='center')
    
    # Datos
    total_general = 0
    for row, pedido in enumerate(pedidos, 2):
        mesa = Mesa.query.get(pedido.mesa_id)
        ws.cell(row=row, column=1, value=pedido.fecha_pago.strftime('%d/%m/%Y'))
        ws.cell(row=row, column=2, value=pedido.fecha_pago.strftime('%H:%M'))
        ws.cell(row=row, column=3, value=f"{mesa.zona} {mesa.numero}")
        ws.cell(row=row, column=4, value=float(pedido.total))
        ws.cell(row=row, column=5, value=pedido.forma_pago.title())
        ws.cell(row=row, column=6, value=pedido.id)
        total_general += float(pedido.total)
    
    # Totales
    row_total = len(pedidos) + 3
    ws.cell(row=row_total, column=3, value="TOTAL:").font = Font(bold=True)
    ws.cell(row=row_total, column=4, value=total_general).font = Font(bold=True)
    
    # Ajustar ancho de columnas
    for col in ws.columns:
        max_length = max(len(str(cell.value)) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = max_length + 2
    
    # Guardar en memoria
    output = BytesIO()
    wb.save(output)
    output.seek(0)
    
    from flask import send_file
    return send_file(
        output,
        as_attachment=True,
        download_name=f'ventas_{periodo}_{today.strftime("%Y%m%d")}.xlsx',
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

@app.route('/export_word/<periodo>')
def export_word(periodo):
    """Export sales report to Word"""
    today = date.today()
    
    if periodo == 'dia':
        fecha_inicio = today
        fecha_fin = today
        titulo = f"Reporte de Ventas del Día - {today.strftime('%d/%m/%Y')}"
    elif periodo == 'semana':
        fecha_inicio = today - timedelta(days=today.weekday())
        fecha_fin = today
        titulo = f"Reporte de Ventas de la Semana - {fecha_inicio.strftime('%d/%m')} al {fecha_fin.strftime('%d/%m/%Y')}"
    else:  # mes
        fecha_inicio = today.replace(day=1)
        fecha_fin = today
        titulo = f"Reporte de Ventas del Mes - {fecha_inicio.strftime('%B %Y')}"
    
    # Obtener datos
    pedidos = Pedido.query.filter(
        Pedido.fecha_pago >= fecha_inicio,
        Pedido.fecha_pago <= fecha_fin + timedelta(days=1),
        Pedido.estado == 'pagado'
    ).all()
    
    # Crear documento
    doc = Document()
    
    # Título
    titulo_p = doc.add_heading('BAR YEDRA', 0)
    titulo_p.alignment = 1  # Centrado
    
    subtitulo = doc.add_heading(titulo, 1)
    subtitulo.alignment = 1
    
    # Estadísticas
    total_efectivo = sum(float(p.total) for p in pedidos if p.forma_pago == 'efectivo')
    total_tarjeta = sum(float(p.total) for p in pedidos if p.forma_pago == 'tarjeta')
    total_general = total_efectivo + total_tarjeta
    
    doc.add_paragraph(f"Total de pedidos: {len(pedidos)}")
    doc.add_paragraph(f"Total en efectivo: €{total_efectivo:.2f}")
    doc.add_paragraph(f"Total en tarjeta: €{total_tarjeta:.2f}")
    doc.add_paragraph(f"TOTAL GENERAL: €{total_general:.2f}")
    
    # Tabla de ventas
    if pedidos:
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        # Encabezados
        encabezados = ['Fecha', 'Hora', 'Mesa', 'Total', 'Forma de Pago', 'Ticket #']
        for i, encabezado in enumerate(encabezados):
            table.cell(0, i).text = encabezado
        
        # Datos
        for pedido in pedidos:
            mesa = Mesa.query.get(pedido.mesa_id)
            row_cells = table.add_row().cells
            row_cells[0].text = pedido.fecha_pago.strftime('%d/%m/%Y')
            row_cells[1].text = pedido.fecha_pago.strftime('%H:%M')
            row_cells[2].text = f"{mesa.zona} {mesa.numero}"
            row_cells[3].text = f"€{pedido.total:.2f}"
            row_cells[4].text = pedido.forma_pago.title()
            row_cells[5].text = str(pedido.id)
    
    # Guardar en memoria
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    from flask import send_file
    return send_file(
        output,
        as_attachment=True,
        download_name=f'ventas_{periodo}_{today.strftime("%Y%m%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route('/export_pdf/<periodo>')
def export_pdf(periodo):
    """Export sales report to PDF"""
    today = date.today()
    
    if periodo == 'dia':
        fecha_inicio = today
        fecha_fin = today
        titulo = f"Ventas del Día - {today.strftime('%d/%m/%Y')}"
    elif periodo == 'semana':
        fecha_inicio = today - timedelta(days=today.weekday())
        fecha_fin = today
        titulo = f"Ventas de la Semana - {fecha_inicio.strftime('%d/%m')} al {fecha_fin.strftime('%d/%m/%Y')}"
    else:  # mes
        fecha_inicio = today.replace(day=1)
        fecha_fin = today
        titulo = f"Ventas del Mes - {fecha_inicio.strftime('%B %Y')}"
    
    # Obtener datos
    pedidos = Pedido.query.filter(
        Pedido.fecha_pago >= fecha_inicio,
        Pedido.fecha_pago <= fecha_fin + timedelta(days=1),
        Pedido.estado == 'pagado'
    ).all()
    
    # Crear PDF
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()
    
    # Título
    title = Paragraph("BAR YEDRA", styles['Title'])
    subtitle = Paragraph(titulo, styles['Heading1'])
    story.append(title)
    story.append(subtitle)
    story.append(Spacer(1, 12))
    
    # Estadísticas
    total_efectivo = sum(float(p.total) for p in pedidos if p.forma_pago == 'efectivo')
    total_tarjeta = sum(float(p.total) for p in pedidos if p.forma_pago == 'tarjeta')
    total_general = total_efectivo + total_tarjeta
    
    stats_data = [
        ['Total de pedidos:', str(len(pedidos))],
        ['Total en efectivo:', f"€{total_efectivo:.2f}"],
        ['Total en tarjeta:', f"€{total_tarjeta:.2f}"],
        ['TOTAL GENERAL:', f"€{total_general:.2f}"]
    ]
    
    stats_table = Table(stats_data)
    stats_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('FONTNAME', (0, 3), (1, 3), 'Helvetica-Bold'),
    ]))
    
    story.append(stats_table)
    story.append(Spacer(1, 20))
    
    # Tabla de ventas
    if pedidos:
        data = [['Fecha', 'Hora', 'Mesa', 'Total', 'Forma de Pago', 'Ticket #']]
        for pedido in pedidos:
            mesa = Mesa.query.get(pedido.mesa_id)
            data.append([
                pedido.fecha_pago.strftime('%d/%m/%Y'),
                pedido.fecha_pago.strftime('%H:%M'),
                f"{mesa.zona} {mesa.numero}",
                f"€{pedido.total:.2f}",
                pedido.forma_pago.title(),
                str(pedido.id)
            ])
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.green),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
    
    doc.build(story)
    output.seek(0)
    
    from flask import send_file
    return send_file(
        output,
        as_attachment=True,
        download_name=f'ventas_{periodo}_{today.strftime("%Y%m%d")}.pdf',
        mimetype='application/pdf'
    )

# Create tables and initialize data
with app.app_context():
    db.create_all()
    
    # Initialize sample data if tables are empty
    if Mesa.query.count() == 0:
        # Create tables for each zone
        zonas = {
            'Terraza': 8,  # 8 tables in terrace
            'Sala': 5,     # 5 tables in main room
            'Barra': 5     # 5 bar spots
        }
        
        for zona, num_mesas in zonas.items():
            for i in range(1, num_mesas + 1):
                mesa = Mesa(numero=i, zona=zona, estado='libre')
                db.session.add(mesa)
        
        # Add sample products with photos
        productos_iniciales = [
            # Cafés y bebidas calientes
            ('Café Solo', 1.20, 'Cafés', 'https://images.unsplash.com/photo-1509042239860-f550ce710b93?w=200&h=200&fit=crop&crop=center'),
            ('Café con Leche', 1.40, 'Cafés', 'https://images.unsplash.com/photo-1461023058943-07fcbe16d735?w=200&h=200&fit=crop&crop=center'),
            ('Cortado', 1.30, 'Cafés', 'https://images.unsplash.com/photo-1495474472287-4d71bcdd2085?w=200&h=200&fit=crop&crop=center'),
            ('Té', 1.20, 'Cafés', 'https://images.unsplash.com/photo-1558618047-3c8c76ca7d13?w=200&h=200&fit=crop&crop=center'),
            ('Carajillo', 4.10, 'Licores', 'https://images.unsplash.com/photo-1517701604599-bb29b565090c?w=200&h=200&fit=crop&crop=center'),
            
            # Bebidas frías
            ('Agua', 1.50, 'Bebidas', 'https://images.unsplash.com/photo-1560472354-b33ff0c44a43?w=200&h=200&fit=crop&crop=center'),
            ('Coca Cola', 2.00, 'Bebidas', 'https://images.unsplash.com/photo-1629203851122-3726ecdf080e?w=200&h=200&fit=crop&crop=center'),
            
            # Alcoholes
            ('Cerveza', 2.50, 'Cervezas', 'https://images.unsplash.com/photo-1618885472179-5e474019f2a9?w=200&h=200&fit=crop&crop=center'),
            ('Vino Tinto', 2.00, 'Vinos', 'https://images.unsplash.com/photo-1506377247377-2a5b3b417ebb?w=200&h=200&fit=crop&crop=center'),
            ('Vino Blanco', 2.00, 'Vinos', 'https://images.unsplash.com/photo-1547595628-c61a29f496f0?w=200&h=200&fit=crop&crop=center'),
            
            # Licores
            ('Baileys', 4.10, 'Licores', 'https://images.unsplash.com/photo-1514362545857-3bc16c4c7d1b?w=200&h=200&fit=crop&crop=center'),
            ('Licores sin alcohol', 2.90, 'Licores', 'https://images.unsplash.com/photo-1551538827-9c037cb4f32a?w=200&h=200&fit=crop&crop=center'),
            ('Orujo, hierbas, pacharán', 2.90, 'Licores', 'https://images.unsplash.com/photo-1569529465841-dfecdab7503b?w=200&h=200&fit=crop&crop=center'),
            ('Anís, brandi o ponche', 2.70, 'Licores', 'https://images.unsplash.com/photo-1569529465841-dfecdab7503b?w=200&h=200&fit=crop&crop=center'),
            ('Jägermeister', 4.10, 'Licores', 'https://images.unsplash.com/photo-1572490122747-3968b75cc699?w=200&h=200&fit=crop&crop=center'),
            ('Cardhu', 8.50, 'Licores', 'https://images.unsplash.com/photo-1569529465841-dfecdab7503b?w=200&h=200&fit=crop&crop=center'),
            ('Whisky (Jony, JB, Ballantines)', 4.10, 'Licores', 'https://images.unsplash.com/photo-1527281400683-1aae777175f8?w=200&h=200&fit=crop&crop=center'),
            ('Chupito nacional', 1.70, 'Licores', 'https://images.unsplash.com/photo-1551538827-9c037cb4f32a?w=200&h=200&fit=crop&crop=center'),
            ('Chupito de anís, brandi, whisky', 1.00, 'Licores', 'https://images.unsplash.com/photo-1551538827-9c037cb4f32a?w=200&h=200&fit=crop&crop=center'),
            ('MAGNo, DYC, Larios, Vodka', 3.10, 'Licores', 'https://images.unsplash.com/photo-1560006082-5324937ce827?w=200&h=200&fit=crop&crop=center'),
            
            # Licores Premium y Combinados
            ('DYC combinado nacional', 5.10, 'Licores Premium', 'https://images.unsplash.com/photo-1570197788417-0e82375c9371?w=200&h=200&fit=crop&crop=center'),
            ('Combinado importación', 6.10, 'Licores Premium', 'https://images.unsplash.com/photo-1564391999744-7ad9e5340204?w=200&h=200&fit=crop&crop=center'),
            ('Bacardí Refresco', 5.50, 'Licores Premium', 'https://images.unsplash.com/photo-1598300042247-d088f8ab3a91?w=200&h=200&fit=crop&crop=center'),
            ('Vodka Refresco', 5.50, 'Licores Premium', 'https://images.unsplash.com/photo-1514362545857-3bc16c4c7d1b?w=200&h=200&fit=crop&crop=center'),
            ('Ginebras Premium', 8.90, 'Licores Premium', 'https://images.unsplash.com/photo-1551538827-9c037cb4f32a?w=200&h=200&fit=crop&crop=center'),
            ('Rones Premium', 6.90, 'Licores Premium', 'https://images.unsplash.com/photo-1598300042247-d088f8ab3a91?w=200&h=200&fit=crop&crop=center'),
            ('Cardhu', 7.90, 'Licores Premium', 'https://images.unsplash.com/photo-1527281400683-1aae777175f8?w=200&h=200&fit=crop&crop=center'),
            
            # Comidas
            ('Tostada con Tomate', 2.50, 'Comidas', 'https://images.unsplash.com/photo-1565299624946-b28f40a0ca4b?w=200&h=200&fit=crop&crop=center'),
            ('Bocadillo Jamón', 4.50, 'Comidas', 'https://images.unsplash.com/photo-1539252554453-80ab65ce3586?w=200&h=200&fit=crop&crop=center'),
            ('Tortilla Española', 3.50, 'Comidas', 'https://images.unsplash.com/photo-1565299624946-b28f40a0ca4b?w=200&h=200&fit=crop&crop=center'),
            ('Patatas Bravas', 3.00, 'Comidas', 'https://images.unsplash.com/photo-1573821663912-6df460f9c684?w=200&h=200&fit=crop&crop=center'),
            ('Aceitunas', 2.00, 'Comidas', 'https://images.unsplash.com/photo-1452827073306-6e6e661baf57?w=200&h=200&fit=crop&crop=center'),
        ]
        
        for nombre, precio, categoria, foto_url in productos_iniciales:
            producto = Producto(
                nombre=nombre, 
                precio=precio, 
                categoria=categoria,
                foto_url=foto_url,
                stock_actual=25,
                stock_minimo=5,
                unidad_medida='unidad'
            )
            db.session.add(producto)
        
        db.session.commit()
        print("Database initialized with sample data")

@app.route('/remove_from_order', methods=['POST'])
def remove_from_order():
    """Remove product from current order"""
    mesa_id = request.form.get('mesa_id')
    producto_id = int(request.form.get('producto_id'))
    
    # Get current order
    pedido = Pedido.query.filter_by(mesa_id=mesa_id, estado='abierto').first()
    if not pedido:
        flash('No hay pedido activo para esta mesa', 'error')
        return redirect(url_for('table_detail', mesa_id=mesa_id))
    
    # Remove product from JSON (using new format: {product_id: quantity})
    import json
    productos_dict = json.loads(pedido.productos) if pedido.productos else {}
    producto_id_str = str(producto_id)
    
    if producto_id_str in productos_dict:
        del productos_dict[producto_id_str]
    
    if len(productos_dict) == 0:
        # If no products left, delete the order and set table as free
        mesa = Mesa.query.get(mesa_id)
        if mesa:
            mesa.estado = 'libre'
        db.session.delete(pedido)
        db.session.commit()
        flash('Producto eliminado. Mesa liberada al no quedar productos', 'info')
    else:
        # Recalculate total with terraza supplement if applicable
        mesa = Mesa.query.get(mesa_id)
        total = 0
        for pid, qty in productos_dict.items():
            prod = Producto.query.get(int(pid))
            if prod:
                precio_unitario = float(prod.precio)
                if mesa.zona.lower() == 'terraza' and not getattr(prod, 'sin_suplemento_terraza', False):
                    precio_unitario += 0.20
                total += precio_unitario * qty
        
        pedido.productos = json.dumps(productos_dict)
        pedido.total = total
        db.session.commit()
        flash('Producto eliminado del pedido', 'success')
    
    return redirect(url_for('table_detail', mesa_id=mesa_id))

@app.route('/split_bill', methods=['POST'])
def split_bill():
    """Calculate split bill for a table"""
    mesa_id = request.form.get('mesa_id')
    personas = int(request.form.get('personas', 2))
    
    # Get current order
    pedido = Pedido.query.filter_by(mesa_id=mesa_id, estado='abierto').first()
    if not pedido:
        return jsonify({'success': False, 'error': 'No hay pedido activo'})
    
    if personas <= 1:
        return jsonify({'success': False, 'error': 'Debe ser entre 2 o más personas'})
    
    amount_per_person = float(pedido.total) / personas
    
    return jsonify({
        'success': True,
        'total': float(pedido.total),
        'personas': personas,
        'amount_per_person': round(amount_per_person, 2)
    })

@app.route('/carta')
def carta():
    """Display the complete menu with all products by category"""
    productos = Producto.query.filter_by(activo=True).order_by(Producto.categoria, Producto.nombre).all()
    
    # Organizar productos por categoría
    productos_por_categoria = {}
    for producto in productos:
        if producto.categoria not in productos_por_categoria:
            productos_por_categoria[producto.categoria] = []
        productos_por_categoria[producto.categoria].append(producto)
    
    return render_template('carta.html', productos_por_categoria=productos_por_categoria)

# === SISTEMA DE FICHAJE DE EMPLEADOS ===

@app.route('/empleados')
def empleados():
    """Lista de empleados y gestión"""
    empleados = Empleado.query.filter_by(activo=True).order_by(Empleado.nombre).all()
    return render_template('empleados.html', empleados=empleados)

@app.route('/empleados/nuevo', methods=['GET', 'POST'])
def nuevo_empleado():
    """Crear nuevo empleado"""
    if request.method == 'POST':
        try:
            empleado = Empleado(
                nombre=request.form['nombre'],
                apellidos=request.form['apellidos'],
                dni=request.form['dni'],
                telefono=request.form.get('telefono', ''),
                email=request.form.get('email', ''),
                puesto=request.form['puesto'],
                horas_contrato_semanal=int(request.form.get('horas_contrato_semanal', 30)),
                pin_fichaje=request.form['pin_fichaje']
            )
            
            db.session.add(empleado)
            db.session.commit()
            flash('Empleado creado exitosamente', 'success')
            return redirect(url_for('empleados'))
            
        except Exception as e:
            flash(f'Error al crear empleado: {str(e)}', 'error')
            db.session.rollback()
    
    return render_template('nuevo_empleado.html')

@app.route('/fichaje')
def fichaje():
    """Pantalla principal de fichaje para empleados"""
    return render_template('fichaje.html')

@app.route('/fichaje/entrada', methods=['POST'])
def fichaje_entrada():
    """Registrar entrada de empleado"""
    pin = request.form.get('pin')
    
    if not pin:
        return jsonify({'success': False, 'error': 'PIN requerido'})
    
    empleado = Empleado.query.filter_by(pin_fichaje=pin, activo=True).first()
    if not empleado:
        return jsonify({'success': False, 'error': 'PIN incorrecto'})
    
    # Verificar si ya tiene fichaje de entrada hoy
    hoy = datetime.utcnow().date()
    fichaje_existente = Fichaje.query.filter_by(
        empleado_id=empleado.id,
        fecha=hoy
    ).first()
    
    if fichaje_existente and fichaje_existente.hora_entrada:
        return jsonify({'success': False, 'error': 'Ya has fichado entrada hoy'})
    
    # Crear o actualizar fichaje
    if not fichaje_existente:
        fichaje = Fichaje(
            empleado_id=empleado.id,
            fecha=hoy,
            hora_entrada=datetime.utcnow()
        )
        db.session.add(fichaje)
    else:
        fichaje_existente.hora_entrada = datetime.utcnow()
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'mensaje': f'Entrada registrada para {empleado.nombre_completo}',
        'hora': datetime.utcnow().strftime('%H:%M')
    })

@app.route('/fichaje/salida', methods=['POST'])
def fichaje_salida():
    """Registrar salida de empleado"""
    pin = request.form.get('pin')
    
    if not pin:
        return jsonify({'success': False, 'error': 'PIN requerido'})
    
    empleado = Empleado.query.filter_by(pin_fichaje=pin, activo=True).first()
    if not empleado:
        return jsonify({'success': False, 'error': 'PIN incorrecto'})
    
    # Buscar fichaje de hoy
    hoy = datetime.utcnow().date()
    fichaje = Fichaje.query.filter_by(
        empleado_id=empleado.id,
        fecha=hoy
    ).first()
    
    if not fichaje or not fichaje.hora_entrada:
        return jsonify({'success': False, 'error': 'No hay entrada registrada hoy'})
    
    if fichaje.hora_salida:
        return jsonify({'success': False, 'error': 'Ya has fichado salida hoy'})
    
    # Registrar salida y calcular horas
    fichaje.hora_salida = datetime.utcnow()
    fichaje.calcular_horas()
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'mensaje': f'Salida registrada para {empleado.nombre_completo}',
        'hora': datetime.utcnow().strftime('%H:%M'),
        'horas_trabajadas': float(fichaje.horas_trabajadas) if fichaje.horas_trabajadas else 0
    })

@app.route('/fichajes/informe')
def informe_fichajes():
    """Informe de fichajes por fechas"""
    fecha_inicio = request.args.get('fecha_inicio')
    fecha_fin = request.args.get('fecha_fin')
    
    if not fecha_inicio:
        fecha_inicio = datetime.utcnow().date() - timedelta(days=7)
    else:
        fecha_inicio = datetime.strptime(fecha_inicio, '%Y-%m-%d').date()
        
    if not fecha_fin:
        fecha_fin = datetime.utcnow().date()
    else:
        fecha_fin = datetime.strptime(fecha_fin, '%Y-%m-%d').date()
    
    fichajes = Fichaje.query.filter(
        Fichaje.fecha >= fecha_inicio,
        Fichaje.fecha <= fecha_fin
    ).order_by(Fichaje.fecha.desc(), Fichaje.hora_entrada).all()
    
    return render_template('informe_fichajes.html', 
                         fichajes=fichajes,
                         fecha_inicio=fecha_inicio,
                         fecha_fin=fecha_fin)

# Create database tables
with app.app_context():
    db.create_all()

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
